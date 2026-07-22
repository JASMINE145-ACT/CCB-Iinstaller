"""Tests for WANd DocumentSpec pipeline."""

from __future__ import annotations

import json
import os
import tempfile
import unittest

from wand_document_spec.gate_r import validate_rendered_document
from wand_document_spec.gate_s import ValidationError, validate_document_spec
from wand_document_spec.patch import patch_block_by_id
from wand_document_spec.render import render_document_spec

FIXTURES = os.path.join(os.path.dirname(__file__), "fixtures")


def _load(name: str) -> dict:
    with open(os.path.join(FIXTURES, name), encoding="utf-8") as f:
        return json.load(f)


class TestGateS(unittest.TestCase):
    def test_valid_spec_passes(self) -> None:
        result = validate_document_spec(_load("valid-report.json"))
        self.assertEqual(result["status"], "PASS")
        self.assertEqual(result["block_count"], 2)

    def test_duplicate_block_id_fails(self) -> None:
        with self.assertRaises(ValidationError):
            validate_document_spec(_load("invalid-duplicate-id.json"))

    def test_placeholder_fails(self) -> None:
        spec = _load("valid-report.json")
        spec["sections"][0]["blocks"][0]["text"] = "Hello {{name}}"
        with self.assertRaises(ValidationError):
            validate_document_spec(spec)


class TestRenderPipeline(unittest.TestCase):
    def test_render_validate_patch(self) -> None:
        spec = _load("valid-report.json")
        with tempfile.TemporaryDirectory() as tmp:
            docx = os.path.join(tmp, "report.docx")
            env = {
                "apply_id": "apply-test-001",
                "document_id": spec["document_id"],
                "mode": "upsert",
                "backup_before": True,
            }
            r1 = render_document_spec(spec, docx, env)
            self.assertEqual(r1["status"], "PASS")
            self.assertTrue(os.path.isfile(r1["manifest_path"]))

            gate_r = validate_rendered_document(spec, docx)
            self.assertEqual(gate_r["status"], "PASS")

            # idempotent retry
            r2 = render_document_spec(spec, docx, env)
            self.assertEqual(r2["status"], "PASS")
            self.assertEqual(r1["manifest_path"], r2["manifest_path"])

            patched = patch_block_by_id(
                docx,
                "blk-exec-p1",
                {"type": "paragraph", "text": "更新后的摘要段落。"},
                {"apply_id": "apply-patch-001", "document_id": spec["document_id"]},
            )
            self.assertEqual(patched["status"], "PASS")

            gate_r2 = validate_rendered_document(spec, docx)
            self.assertEqual(gate_r2["status"], "PASS")

    def test_restore_backup(self) -> None:
        spec = _load("valid-report.json")
        with tempfile.TemporaryDirectory() as tmp:
            docx = os.path.join(tmp, "report.docx")
            env = {
                "apply_id": "apply-restore-001",
                "document_id": spec["document_id"],
                "mode": "upsert",
                "backup_before": True,
            }
            render_document_spec(spec, docx, env)
            from docx import Document

            Document(docx).add_paragraph("corrupt")
            Document(docx).save(docx)

            from wand_document_spec.apply import restore_document_backup

            restored = restore_document_backup(docx, "apply-restore-001")
            self.assertEqual(restored["status"], "PASS")
            gate_r = validate_rendered_document(spec, docx)
            self.assertEqual(gate_r["status"], "PASS")


if __name__ == "__main__":
    unittest.main()
