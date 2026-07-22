"""Render DocumentSpec to DOCX with bookmarks and manifest."""

from __future__ import annotations

import hashlib
import json
import os
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from wand_document_spec.apply import (
    ApplyEnvelope,
    backup_document,
    prepare_mutate,
    record_apply,
    spec_hash,
)
from wand_document_spec.gate_s import validate_document_spec
from wand_document_spec.models import Block, DocumentSpec, Section, parse_spec

_BOOKMARK_SEQ = 1


def _next_bookmark_id() -> int:
    global _BOOKMARK_SEQ
    bid = _BOOKMARK_SEQ
    _BOOKMARK_SEQ += 1
    return bid


def _bookmark_name(kind: str, entity_id: str) -> str:
    return f"WANd:{kind}:{entity_id}"


def _add_bookmark(paragraph, name: str) -> int:
    bid = _next_bookmark_id()
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    p_el = paragraph._p
    insert_at = 1 if p_el.find(qn("w:pPr")) is not None else 0
    p_el.insert(insert_at, start)
    p_el.append(end)
    return bid


def _content_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()[:16]


def manifest_path_for(docx_path: str) -> str:
    return f"{docx_path}.manifest.json"


def _heading_style(level: int) -> str:
    return f"Heading {min(max(level, 1), 3)}"


def _render_block(doc: Document, block: Block, manifest_entries: list[dict[str, Any]]) -> None:
    if block.type == "page_break":
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        bname = _bookmark_name("blk", block.block_id)
        _add_bookmark(p, bname)
        manifest_entries.append(
            {
                "block_id": block.block_id,
                "bookmark_name": bname,
                "type": block.type,
                "content_hash": _content_hash("page_break"),
            }
        )
        return

    if block.type == "heading":
        level = block.level or 2
        p = doc.add_paragraph(block.text or "", style=_heading_style(level))
        bname = _bookmark_name("blk", block.block_id)
        _add_bookmark(p, bname)
        manifest_entries.append(
            {
                "block_id": block.block_id,
                "bookmark_name": bname,
                "type": block.type,
                "content_hash": _content_hash(block.text or ""),
            }
        )
        return

    if block.type == "paragraph":
        p = doc.add_paragraph(block.text or "")
        bname = _bookmark_name("blk", block.block_id)
        _add_bookmark(p, bname)
        manifest_entries.append(
            {
                "block_id": block.block_id,
                "bookmark_name": bname,
                "type": block.type,
                "content_hash": _content_hash(block.text or ""),
            }
        )
        return

    if block.type == "list":
        style = "List Number" if block.ordered else "List Bullet"
        for item in block.items or []:
            p = doc.add_paragraph(item, style=style)
        # single bookmark on first list item paragraph
        if doc.paragraphs:
            p = doc.paragraphs[-len(block.items or [])]
            bname = _bookmark_name("blk", block.block_id)
            _add_bookmark(p, bname)
            manifest_entries.append(
                {
                    "block_id": block.block_id,
                    "bookmark_name": bname,
                    "type": block.type,
                    "content_hash": _content_hash("|".join(block.items or [])),
                }
            )
        return

    if block.type == "table":
        headers = block.headers or []
        rows = block.rows or []
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        for ci, h in enumerate(headers):
            table.rows[0].cells[ci].text = h
        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row):
                table.rows[ri + 1].cells[ci].text = cell
        # bookmark on paragraph after table (python-docx inserts implicit p)
        p = doc.add_paragraph()
        p.add_run("")  # anchor paragraph
        bname = _bookmark_name("blk", block.block_id)
        _add_bookmark(p, bname)
        flat = "\n".join(["|".join(headers)] + ["|".join(r) for r in rows])
        manifest_entries.append(
            {
                "block_id": block.block_id,
                "bookmark_name": bname,
                "type": block.type,
                "content_hash": _content_hash(flat),
            }
        )


def _render_section(doc: Document, section: Section, manifest_entries: list[dict[str, Any]]) -> None:
    hp = doc.add_paragraph(section.title, style=_heading_style(section.level))
    sec_bname = _bookmark_name("sec", section.section_id)
    _add_bookmark(hp, sec_bname)
    manifest_entries.append(
        {
            "section_id": section.section_id,
            "bookmark_name": sec_bname,
            "type": "section",
            "title": section.title,
        }
    )
    for block in section.blocks:
        _render_block(doc, block, manifest_entries)


def _write_manifest(
    docx_path: str,
    document_id: str,
    spec: DocumentSpec,
    entries: list[dict[str, Any]],
) -> str:
    mpath = manifest_path_for(docx_path)
    payload = {
        "document_id": document_id,
        "docx_path": os.path.abspath(docx_path),
        "spec_version": spec.spec_version,
        "title": spec.title,
        "entries": entries,
    }
    with open(mpath, "w", encoding="utf-8") as f:
        json.dump(payload, f, indent=2, ensure_ascii=False)
    return mpath


def render_document_spec(
    spec: dict[str, Any] | str,
    filename: str,
    apply_envelope: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """
    Render DocumentSpec to DOCX. Gate S runs first.
    Returns manifest path, entry count, apply report path.
    """
    global _BOOKMARK_SEQ
    _BOOKMARK_SEQ = 1

    raw = json.loads(spec) if isinstance(spec, str) else spec
    gate = validate_document_spec(raw)
    doc_spec = parse_spec(raw)
    envelope = ApplyEnvelope.from_dict(apply_envelope)
    apply_id = envelope.apply_id if envelope else f"render-{doc_spec.document_id}"
    sh = spec_hash(raw)

    rec_path = os.path.join(
        os.path.dirname(os.path.abspath(filename)) or ".",
        ".wand-apply",
        f"{apply_id}.json",
    )
    if os.path.isfile(rec_path) and envelope and envelope.mode == "upsert":
        with open(rec_path, encoding="utf-8") as f:
            rec = json.load(f)
        cached = rec.get("result") if rec.get("status") == "PASS" else None
        if cached and cached.get("spec_hash") == sh:
            return cached

    prepare_mutate(filename, envelope)

    document = Document()
    # title as document heading
    title_p = document.add_paragraph(doc_spec.title, style="Title")
    _add_bookmark(title_p, _bookmark_name("doc", doc_spec.document_id))

    entries: list[dict[str, Any]] = []
    for section in doc_spec.sections:
        _render_section(document, section, entries)

    os.makedirs(os.path.dirname(os.path.abspath(filename)) or ".", exist_ok=True)
    document.save(filename)
    if envelope and envelope.backup_before:
        backup_document(filename, envelope.apply_id)
    mpath = _write_manifest(filename, doc_spec.document_id, doc_spec, entries)

    result = {
        "status": "PASS",
        "gate_s": gate,
        "document_id": doc_spec.document_id,
        "filename": os.path.abspath(filename),
        "manifest_path": mpath,
        "entry_count": len(entries),
        "section_count": gate["section_count"],
        "block_count": gate["block_count"],
        "spec_hash": sh,
    }
    record_apply(filename, apply_id, envelope, "PASS", result, mpath)
    return result


def get_document_manifest(filename: str) -> dict[str, Any]:
    mpath = manifest_path_for(filename)
    if not os.path.isfile(mpath):
        return {"status": "FAIL", "error": f"manifest not found: {mpath}"}
    with open(mpath, encoding="utf-8") as f:
        data = json.load(f)
    data["status"] = "PASS"
    data["manifest_path"] = mpath
    return data
